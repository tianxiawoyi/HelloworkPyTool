#在线选房
import docx
import xlwt

docxPath2=r'D:\act选房namecn.docx'
# docxPath='D:\hjzb数据库字典_简化.docx'
#docxPath=r'D:\act选房.docx'    # r代表非转义
docxPath=r'D:\act在线选房全.docx'
#docxPath=r'D:\华发在线售楼处.doc'

row = 0       # 记录行数
tableSum=0   # 记录表个数
workbook = xlwt.Workbook(encoding='utf-8')
worksheet = workbook.add_sheet("sheet1")
style = xlwt.easyxf('pattern: pattern solid, fore_colour ice_blue')

# 创建一个已存在的 word 文档的对象
file2 = docx.Document(docxPath2)

# 创建一个已存在的 word 文档的对象
file = docx.Document(docxPath)


# 读取每个段落的内容并输出
# for it in file.paragraphs:
#     print( it.text )

paragraphsList = [it.text for it in file.paragraphs if len(it.text)>0]
print("段落数量:"+str(len(file.paragraphs)))
print("表格数量:"+str(len(file.tables)))   #480
print("paragraphsList数量:"+str(len(paragraphsList)))

dist_table={}
# 读取表格中的内容并输出
for index,it in enumerate(file2.tables):
    for i,r in enumerate(it.rows):   # 每行
        cells = r.cells
        dist_table[cells[0].text]= cells[1].text
        # print(cells[0].text)
        # print(cells[1].text)
print(dist_table)
print(len(dist_table))

# 读取每个段落的内容并输出
# for text in paragraphsList:
#     print(text)

# 读取表格中的内容并输出
for index,it in enumerate(file.tables):
    textTable=paragraphsList[index]
    tablename = textTable
    tablename_cn =''
    if tablename in dist_table.keys(): # if dist_table[tablename]  为空 , 取空
        tablename_cn = dist_table[tablename]
    print("----tablename_cn------->"+tablename_cn)
    # splits = textTable.split(':')
    # tablename = splits[1]
    # tablename_cn = splits[2]
    print(tablename)
    if index%2!=0:
        for i,r in enumerate(it.rows):
            if i==0:continue
            row= row+1
            worksheet.write(row, 2, label= index+1 ,style = style)    # 表序号
            worksheet.write(row, 3, label= tablename ,style = style)
            worksheet.write(row, 4, label= tablename_cn ,style = style )
            for j,cell in enumerate(r.cells):
                # print( cell.text )
                worksheet.write(row, j + 5, label= cell.text ,style = style)
    else:
        for i,r in enumerate(it.rows):
            if i==0:continue
            row= row+1
            worksheet.write(row, 2, label= index+1 )   # 表序号
            worksheet.write(row, 3, label= tablename )
            worksheet.write(row, 4, label= tablename_cn )
            for j,cell in enumerate(r.cells):
                # print( cell.text )
                worksheet.write(row, j + 5, label= cell.text )


 # print(tableSum)

print("段落数量:"+str(len(file.paragraphs)))
print("表格数量:"+str(len(file.tables)))   #513
print("paragraphsList数量:"+str(len(paragraphsList)))

xlsSavePath = "xls/在线选房/在线选房8.xls"
workbook.save(xlsSavePath)