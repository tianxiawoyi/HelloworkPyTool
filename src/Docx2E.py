import docx
import xlwt

# docxPath='D:\hjzb数据库字典_简化.docx'
#docxPath=r'D:\act选房.docx'    # r代表非转义
docxPath=r'D:\act选房namecn.docx'
#docxPath=r'D:\华发在线售楼处.doc'

row = 0       # 记录行数
tableSum=0   # 记录表个数
workbook = xlwt.Workbook(encoding='utf-8')
worksheet = workbook.add_sheet("sheet1")
style = xlwt.easyxf('pattern: pattern solid, fore_colour ice_blue')

# 创建一个已存在的 word 文档的对象
file = docx.Document(docxPath)


# 读取每个段落的内容并输出
# for it in file.paragraphs:
#     print( it.text )

paragraphsList = [it.text for it in file.paragraphs if len(it.text)>0]
print("段落数量:"+str(len(file.paragraphs)))
print("表格数量:"+str(len(file.tables)))   #480
print("paragraphsList数量:"+str(len(paragraphsList)))

# 读取每个段落的内容并输出
for text in paragraphsList:
    print(text)

dist_table={}
# 读取表格中的内容并输出
for index,it in enumerate(file.tables):
    for i,r in enumerate(it.rows):   # 每行
        cells = r.cells
        dist_table[cells[0].text]= cells[1].text
        print(cells[0].text)
        print(cells[1].text)
print(dist_table)
print(len(dist_table))